import time
import win32com.client
from docx import Document
import re
import os

wordapp = win32com.client.Dispatch("Word.Application") 
wordapp.Visible = 1


#worddoc.Content.MoveEnd
# Close the Word Document (a save-Dialog pops up)
#worddoc.Close()
# Close the Word Application
#wordapp.Quit()

# https://python-docx.readthedocs.io/en/latest/user/documents.html

# https://buildmedia.readthedocs.org/media/pdf/python-docx/latest/python-docx.pdf
# https://relopezbriega.github.io/blog/2015/07/19/expresiones-regulares-con-python/

#filename = '308465'
#filename = '356535'
#doc = Document("./documents/"+filename+".docx")

print(os.getcwd())



ejemplo_dir = os.getcwd()+'\\readword\\documents\\'
contenido = os.listdir(ejemplo_dir)
print("Documentos analizar",contenido)

imagenes = []
for fichero in contenido:
    if os.path.isfile(os.path.join(ejemplo_dir, fichero)) and fichero.endswith('.docx'):
         print (fichero)
         filename = os.path.splitext(fichero)[0]
         #filename = fichero[0:6]
         rpsta = os.path.exists(ejemplo_dir+filename+".docx")
         if rpsta:
            print("SI Existe Documento")
         else:
            print("NO Existe Documento",ejemplo_dir+filename+".docx")

         worddoc = wordapp.Documents.Add()
         worddoc.Content.Font.Size = 12
         doc = Document(ejemplo_dir+filename+".docx")



         print("verifica contenido",doc)
         fullText = []
         for para in doc.paragraphs:
            fullText.append(para.text)
         text = '\n'.join(fullText)

         

         reg_kardex = 'K[0-9]{5},'
         reg_numero = ' [0-9]{8},'
         reg_detalle = 'DOCUMENTO NACIONAL DE IDENTIDAD NUMERO [0-9]+'
         reg_partida = 'PARTIDA N° [0-9]+'
         reg_partida2 = 'PARTIDA N° P[0-9]+'

         time.sleep(3)

         worddoc.Content.InsertAfter("Reporte de Incidencias KARDEX N°"+filename +" \n")
         worddoc.Content.InsertAfter("\n")

         print("Numero KARDEX------"+ filename)

         print("Numero DNI------")
         a=re.findall(reg_numero,text)
         worddoc.Content.InsertAfter("1. Numero DNI"+ str(a) +" \n")
         print(a)
         print("Detalle DNI------")
         b=re.findall(reg_detalle,text)
         worddoc.Content.InsertAfter("2. Detalle "+ str(b) +" \n")
         print(b)

         print("Partida Registral------")
         b=re.findall(reg_partida,text)
         worddoc.Content.InsertAfter("3. Partida Registral "+ str(b) +" \n")
         print("Partida Registral" +str(b))
         if not b:
            print ("Consultar por el Nombre existe una P")
            worddoc.Content.InsertAfter("Consultar por el codigo de Partida existe una 'P' \n")
            b=re.findall(reg_partida2,text)
            print(b)


         worddoc.Content.InsertAfter("4. Verificar DNI "+" \n")

         worddoc.Content.InsertAfter("5. Verificar Partida Registral "+" \n")

         worddoc.Content.InsertAfter("6. Verificar RUC "+" \n")

         worddoc.Content.InsertAfter("7. Verificar Verificación Prico y Declaraciones "+" \n\r")


         time.sleep(1)

         # todos los numeros
         a=re.findall(r'\d+',text)
         print (a)
         worddoc.Content.InsertAfter("Resumen de numero obtenidos: \n"+str(a) + " \n\r")





