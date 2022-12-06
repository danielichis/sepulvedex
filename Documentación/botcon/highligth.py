

from docx.enum.text import WD_COLOR_INDEX
from docx import Document

#Instalarte ademas este modulo para los comments
#pip install bayoo-docx

kardex_query = '376150'
xruta='C:\\Datafiles\\'
# documento
template_file_path = xruta+kardex_query+'.docx'

def writedoc():
    print('inserta Anotaciones')
    template_document = Document(template_file_path)

    for paragraph in template_document.paragraphs:
        if 'PRIMA FITNESS S.A.C.' in paragraph.text:
            for run in paragraph.runs:
                if 'PRIMA FITNESS S.A.C.' in run.text:
                    x = run.text.split('PRIMA FITNESS S.A.C.')
                    run.clear()
                    for i in range(len(x)-1):
                        run.add_text(x[i])
                        run.add_text('PRIMA FITNESS S.A.C.')
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW




    paragraph = template_document.add_paragraph('EJEMPLO DE OBSERVACION') # create new paragraph

    comment = paragraph.add_comment('Comentario u Observacion2',author='Bot Confrontacion',initials= 'bc') # add a comment on the entire paragraph

    paragraph2 = template_document.add_paragraph('Comentario u Observacion') # create another paragraph

    run = paragraph2.add_run('Comentario u Observacion2') #add a run to the paragraph


    run.add_comment('Otra Observacion Encontrada',author='Observacion de DNI') # add a comment only for the run text 

    run_comments = run.comments

    paragraph.add_footnote('footnote text') # add a footnote
    print("Documento Verificado")
    template_document.save(xruta+'Kardex_verificado'+kardex_query+'.docx')