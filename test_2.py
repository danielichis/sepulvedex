import docx


doc = docx.Document('./KardexsOut/K42775.docx')

# for p in doc.paragraphs:
#     for run in p.runs:
#         print(run.text)

for p in doc.paragraphs:
    if 'graphicData' in p._p.xml:
        print(p._p.xml)
        print('grp')

# doc.paragraphs[0].runs[0].add_comment('img presente en el documento')

# doc.save('./KardexsOut/K45970.docx')