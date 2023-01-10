import docx
import pandas as pd
import re
import math
import os
from utilities import pathsManager as pm
print(pm().currentFolderPath)

def amountsValidation(docxPathOut):
    doc = docx.Document(docxPathOut)
    for para in doc.paragraphs:
        if para.text.strip() == "MUTUO CON GARANTIA HIPOTECARIA":
            cellAmounts = []
            rowText = []
            tableText = []
            tables = {}
            wholeTable = {}
            rowsLens = []
            realDf = pd.DataFrame()

            for i, table in enumerate(doc.tables):
                for row in table.rows:
                    for j,cell in enumerate(row.cells):
                        if bool(cell.text) == False:
                            cell.text = "-"
                        columnKey = "Column " + str(j)
                        tableKey = "Table " + str(i)
                        tables.setdefault(tableKey, {})
                        tables[tableKey].setdefault(columnKey, []).append(cell.text)
                        
            # print(tables)
            for d in tables:    
                columnLens = []
                for i in tables[d]:
                    columnLens.append(len(tables[d][i]))

                maxColumnLen = max(columnLens)
                for i in tables[d]:
                    while len(tables[d][i]) < maxColumnLen:
                        tables[d][i].insert(0, "-")

            for k in tables:
                df = pd.DataFrame(tables[k])
                if df.iloc[0,0].strip() == "MONTO DEL PRESTAMO":
                    realDf = df
                    # print(tables[k])
                    break
            # print(realDf.shape)
            # print(realDf)
            for row in range(8, realDf.shape[0]):
                cellAmount = re.findall(r'[s/$]\s*([\d,.]+)', realDf.iloc[row, 0]) 
                cellAmount = cellAmount[0]
                cellAmount = cellAmount.replace(" ", "")
                cellAmount = cellAmount.replace(",", "")
                cellAmounts.append(float(cellAmount))
            

            c = realDf.iloc[3, 4]
            bigAmount = re.findall(r'[\d,.]+', c)
            bigAmount = bigAmount[0]
            bigAmount = bigAmount.replace(" ", "")
            bigAmount = bigAmount.replace(",", "")
            bigAmount = float(bigAmount)

            a = math.floor(sum(cellAmounts))
            b = math.floor(bigAmount)
            
            if a == b:
                para.runs[-1].add_comment("LOS MONTOS COINCIDEN", author='BOT CONFRONT')
                print("The amounts match")
            else:
                para.runs[-1].add_comment("LOS MONTOS NO COINCIDEN, REVISAR MANUALMENTE", author='BOT CONFRONT')
                print("The amounts don't match")
        

            break

    doc.save(docxPathOut)


def amountsValidation0():
    directory = os.path.join(pm().currentFolderPath, 'Kardexs')
    directoryOut = os.path.join(pm().currentFolderPath, 'KardexsOut')
    files = os.listdir(directory)
    files = [f for f in files if f.endswith('.docx')]
    print(files)
    for file in files:
        docxPath = os.path.join(directory, file)
        docxPathOut = os.path.join(directoryOut, file)
        amountsValidation(docxPathOut)

if __name__=='__main__':
    amountsValidation0()

