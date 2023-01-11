from docx.enum.text import WD_COLOR_INDEX
from docx import Document
from utilities import pathsManager
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from numerToLetters import numero_a_moneda_sunat
import re
import os
import json
#Instalarte ademas este modulo para los comments
#pip install bayoo-docx
def split_text(text, word):
    pattern = re.compile(r'([\S\s]*)(\b{})([\S\s]*)'.format(word))
    match = pattern.search(text)
    if match:
        return match.groups()
    return None


def split_Runs(doc,word):
    for p in doc.paragraphs:
        if p.text.find(word) != -1:            
            virtualRuns=p.runs
            p.text = ""
            #x=[st for st in doc.styles if st.style_id == '']
            for r in virtualRuns:
                try:
                    font_styles = doc.styles
                    font_charstyle = font_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
                    font_object = font_charstyle.font
                    font_object.size =r.font.size
                    font_object.name = r.font.name
                except:
                    pass
                if r.text.find(word) != -1:
                    before, word, after = split_text(r.text, word)
                    #clone the format of the original run
                    p.add_run(before,style="CommentsStyle")
                    p.add_run("",style="CommentsStyle")
                    p.add_run(word,style="CommentsStyle")
                    p.add_run(after,style="CommentsStyle")
                else:
                    p.add_run(r.text,style="CommentsStyle")
    return doc
def style_Token(doc,word,comment,specialP=False):
    for p in doc.paragraphs:
        for i,r in enumerate(p.runs):
            if p.runs[i].text.find(word) != -1:
                p.runs[i].font.highlight_color = WD_COLOR_INDEX.YELLOW
                if comment:
                    p.runs[i-1].add_comment(f'El registro no coincide con {word} ',author='BOT CONFRONT')
                    #r.add_comment(f'{word} No se encuentra en el documento',author='BOT CONFRONT')
    return doc
def style_Token2(doc,word,comment):
    p=doc.paragraphs[word["indexP"]]    
    if p.text.find(word["correlative"]) != -1:
        p.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
        if comment:
            p.runs[0].add_comment(f'CORRELATIVO O REFERENCIA INCORRECTO',author='BOT CONFRONT')
            #r.add_comment(f'{word} No se encuentra en el documento',author='BOT CONFRONT')
    return doc
def confront(pathkard,doc, numero,descripction,splitMode):
    alltext=getText(doc)
    #print(f"reading {numero} {descripction}")
    if splitMode:
        doc=split_Runs(doc,numero)
    else:
        if isPresent(descripction,alltext):
            doc=style_Token(doc,numero,False)
        else:
            doc=style_Token(doc,numero,True)
    doc.save(pathkard)
def getText(doc):
    fullText=[]
    for para in doc.paragraphs:
        fullText.append(para.text)
    return " ".join(fullText)
def getNumbers(doc):
    fullText=[]
    numerLetters=[]
    for para in doc.paragraphs:
        fullText.append(para.text)
        numletter=re.findall(r"S\/ (\d{1,}.*?\.\d{2}) \((.*?)\)",para.text)
        if len(numletter)>0:
            numerLetters.append(numletter[0])
    return numerLetters
def isPresent(Dname,alltext):
    if alltext.find(Dname) == -1 :
        return False
    else:
        return True
def validateAllData(value,doc,pathkard,splitMode):
    for key2 in value:
        value2=value[key2]
        for dict3 in value2:
            values=list(dict3.values())
            confront(pathkard,doc, values[0],values[1],splitMode)
def validate_amounts(alltext,doc,pathkardOut):
    # Getting more info
    monts=re.findall(r"([0-9,´]*\.\d{2}) \(([\w\s]* [Y-y]? ?00/100)",alltext)
    for mont in monts:
        y=mont[1].replace("CON","Y")
        numero=mont[0].replace(",","").replace("´","")
        x=numero_a_moneda_sunat(float(numero)).replace(" SOLES","")
        x=x.replace("CON","Y")
        if x!=y:
            print(f"ERROR {mont[0]} {mont[1]} {x}")
            doc=split_Runs(doc, mont[0])
            doc=style_Token(doc, mont[0],True)
            doc.save(pathkardOut)
        else:
            print(f"OK {mont[0]} {mont[1]} {x}")

def readJsonPages():
    pm=pathsManager()
    jsonPath=os.path.join(pm.currentFolderPath,"dataofPages.json")
    with open(jsonPath) as json_file:
        data = json.load(json_file)    
    for key in data:
        print(f"reading ------------{key}----------")
        pathkard=os.path.join(pm.currentFolderPath,"Kardexs",key)
        pathkardOut=os.path.join(pm.currentFolderPath,"KardexsOut",key)
        doc=Document(pathkard)
        doc.save(pathkardOut)
        value = data[key]
        validateAllData(value,doc,pathkardOut,True)
        validateAllData(value,doc,pathkardOut,False)
        doc=Document(pathkardOut)
        validate_amounts(getText(doc),doc,pathkardOut)
#readJsonPages()