import docx
import re
from numerToLetters import numero_a_moneda_sunat
import os
from utilities import configData,pathsManager
from validate import style_Token2
def ValiteConsecutiveNums(correlativos):
  grupo_actual = correlativos[0]["index"]
  for i in range(len(correlativos) - 1):
    if correlativos[i]["index"] + 1 != correlativos[i + 1]["index"]:
      # Si el número actual no es el número anterior más uno,
      # verificamos si es el primer número de un nuevo grupo
      if correlativos[i + 1]["index"] == grupo_actual:
        grupo_actual = correlativos[i + 1]["index"]
      else:
        return correlativos[i + 1]
  return True

def getText(doc):
    fullText=[]
    for para in doc.paragraphs:
        fullText.append(para.text)
    return " ".join(fullText)

def ordinals_numbers(doc):
    fullNums1=[]
    fullNums2=[]
    fullNums3=[]
    fullNums4=[]
    corr=[]
    for j,para in enumerate(doc.paragraphs):
        correlativ=re.findall(r"^([1-9]{1,2}).?[-]?\)? ",para.text)
        for corr in correlativ:
            if corr:
                unitDict={
                    "indexP":j,
                    "index":int(corr),
                    "level":1,
                }
                fullNums1.append(unitDict)
        correlativ2=re.findall(r"^[1-9]{1,2}\.(\d{1,2}).",para.text)
        for corr in correlativ2:
            if corr:
                unitDict={
                    "indexP":j,
                    "index":int(corr),
                    "level":2,
                }
                
                fullNums2.append(unitDict)
        correlativ3=re.findall(r"^[1-9]{1,2}\.\d{1,2}\.(\d{1,2}).?",para.text)
        for corr in correlativ3:
            if corr:
                unitDict={
                    "indexP":j,
                    "index":int(corr),
                    "level":3
                }
                fullNums3.append(unitDict)
        correlativ4=re.findall(r"^[1-9]{1,2}\.\d{1,2}\.\d{1,2}\.(\d{1,2}).",para.text)
        for corr in correlativ4:
            if corr:
                unitDict={
                    "indexP":j,
                    "index":int(corr),
                    "level":4
                }
                fullNums4.append(unitDict)
    total=[fullNums1,fullNums2,fullNums3,fullNums4]
    print(total)
    return total
def correlativeLetters(doc,numLeters):
    #numLeters=["PRIMER","SEGUND","TERCER","CUART","QUINT","SEXT","SÉPTIM","OCTAV","NOVEN","DÉCIM"]
    pattern=""
    for i,leter in enumerate(numLeters):
        for j,c in enumerate(leter["correlative"]):
            if j==0:
                pattern0=r"^(%s *" % c
            else:
                pattern0=pattern0+r"%s *" % c
        if i==0:
            pattern=pattern0+r"[AO])[:.]"
        else:
            pattern=pattern+r"|%s[AO])[:.]" % pattern0


    setNumsList=[]
    for i,p in enumerate(doc.paragraphs):
        setNums=re.findall(pattern,p.text)
        for num in setNums:
            for n in num:
                if n:
                    unitDict={
                        "indexP":i,
                        "correlative":n
                    }
                    setNumsList.append(unitDict)
    #print(setNumsList)
    return setNumsList
    
def getIndexCorrelative(wordCorrelatives,ExcelCorrelatives):
    listIndex=[]
    for i,cor in enumerate(wordCorrelatives):
        for k,netCor in enumerate(ExcelCorrelatives):
            for j,c in enumerate(netCor["correlative"]):
                if j==0:
                    pattern0=r"^(%s *" % c
                else:
                    pattern0=pattern0+r"%s *" % c
            if re.match(pattern0+r"[AO])$",cor["correlative"]):
                parIndex={
                    "index":netCor["index"],
                    "correlative":cor["correlative"],
                    "indexP":cor["indexP"]
                }
                listIndex.append(parIndex)
                #break
    #print(listIndex)
    return listIndex

def validateLetters(doc,numLeters):
    wordCorrelatives=correlativeLetters(doc,numLeters)
    indexsTovalidate=getIndexCorrelative(wordCorrelatives,numLeters)
    indexsErrors=True
    if len(indexsTovalidate)>0:
        indexsErrors=ValiteConsecutiveNums(indexsTovalidate)
    if indexsErrors==True:
        print("CORRELATIVOS CARDINALES CORRECTOS")
    else:
        print("CORRELATIVOS CARDINALES INCORRECTOS")
        print(indexsErrors)
        doc=style_Token2(doc,indexsErrors,True)
    return doc
def validateOrdinals(doc):
    ordinalsTovalidate=ordinals_numbers(doc)
    listIndexError=[]
    for level in ordinalsTovalidate:
        indexsError=True
        if len(level)>0:
            indexsError=ValiteConsecutiveNums(level)
        if indexsError==True:
            pass
        else:
            listIndexError.append(indexsError)
            print(indexsError)
    if len(listIndexError)>0:
        print("CORRELATIVOS ORDINALES INCORRECTOS")
        for indexError in listIndexError:
            doc=style_Token2(doc,indexError,True)
        #print(listIndexError)
    else:
        print("CORRELATIVOS ORDINALES CORRECTOS")

    return doc

# A-->1, B-->2, C-->3, etc.

def alfanumeric_nums(doc):
    alfanumsTovalidate=[]
    for i,p in enumerate(doc.paragraphs):
        alfanums=re.findall(r"^([A-z])\.-|^([A-z])\.|^([A-z])\)",p.text)
        for alfas in alfanums:
            for alf in alfas:
                if alf:
                    index=ord(alf)-64
                    unitDict={
                        "indexP":i,
                        "correlative":alf,
                        "index":index
                    }
                    alfanumsTovalidate.append(unitDict)
    print(alfanumsTovalidate)
    return alfanumsTovalidate
def validateAbdc(doc):
    alfaNumericsTovalidate=alfanumeric_nums(doc)
    indexsError=True
    if len(alfaNumericsTovalidate)>0:
        indexsError=ValiteConsecutiveNums(alfaNumericsTovalidate)
    if indexsError==True:
        print("CORRELATIVOS ALFANUMERICOS CORRECTOS")
    else:
        print("CORRELATIVOS ALFANUMERICOS INCORRECTOS")
        print(indexsError)
        doc=style_Token2(doc,indexsError,True)
    return doc

def gettingReferences(doc):
    listReference=[]
    for i,p in enumerate(doc.paragraphs):
        refs=re.findall(r'(\w{5,}) MODIFICACION',p.text)
        if len(refs)>0:
            for ref in refs:
                unitDict={
                    "indexP":i,
                    "correlative":ref,
                }
                listReference.append(unitDict)
    return listReference
def validateReferences(doc):
    referenceTovalidate=gettingReferences(doc)
    if len(referenceTovalidate)>1:
        if referenceTovalidate[0]["correlative"]==referenceTovalidate[1]["correlative"]:
            print("REFERENCIA CORRECTA")
        else:
            print("REFERENCIA INCORRECTA")
            doc=style_Token2(doc,referenceTovalidate[1],True)
    return doc
def validateCorrelatives():
    pm=pathsManager()
    listOfDocxFiles=[f for f in os.listdir(os.path.join(pm.currentFolderPath,"Kardexs")) if f.endswith(".docx")]
    cnfd=configData(os.path.join(pm.currentFolderPath,"config.xlsx"))
    numLeters=cnfd.get_data_config()
    numLetersExcel=numLeters["CORRELATIVOS"].values.tolist()
    numLeters=[]
    for i,corrWords in enumerate(numLetersExcel):
        splitWords=corrWords.split(",")
        for word in splitWords:
            dictCorrelatives={
                "index":i+1,
                "correlative":word}
            numLeters.append(dictCorrelatives)
    
    for file in listOfDocxFiles:
        print(file)
        krdxOutP=os.path.join(pm.currentFolderPath,"KardexsOut",file)
        krdxP=os.path.join(pm.currentFolderPath,"Kardexs",file)
        doc=docx.Document(krdxP)
        doc=validateLetters(doc,numLeters)
        doc=validateOrdinals(doc)
        doc=validateAbdc(doc)
        doc=validateReferences(doc)
        doc.save(krdxOutP)
validateCorrelatives()
#listOfDocxFiles=[f for f in os.listdir(r"C:\DanielBots\Sepulveda\sepulvedex\Kardexs") if f.endswith(".docx")]

# for file in listOfDocxFiles:
#     print(file)
#     doc=docx.Document(r"C:\DanielBots\Sepulveda\sepulvedex\Kardexs\\"+file)
#     text=getalltext(doc)

