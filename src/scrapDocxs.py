import docx
from utilities import configData
from utilities import pathsManager
import json
#from scrapPages import scrapingPages
import os
import re
#get current path
class dataWord:
    def __init__(self,path) -> None:
        pm=pathsManager()
        self.kardexName=path
        self.kardexsPath=os.path.join(pm.currentFolderPath,"Kardexs",path)
        self.configPath=os.path.join(pm.currentFolderPath,"config.xlsx")
        self.doc=docx.Document(self.kardexsPath)
        self.fullText=self.getText()
        self.configData=configData(self.configPath)
        self.dataWord=None

    def getText(self):
        fullText=[]
        for para in self.doc.paragraphs:
            r=para.add_run()
            fullText.append(para.text)
        self.fullText=" ".join(fullText)

    def get_names(self,typeKeys,nameKeys):
        dnislist=[]
        if typeKeys=="DNIS":
            n=8
            midpattern=r" "
            pass
        elif typeKeys=="RUCS":
            n=11
            midpattern=r" "
            pass
        elif typeKeys=="PARTIDAS":
            n=8
            midpattern=r".*?"
            pass
        elif typeKeys=="CES":
            pass
            midpattern=r" "
            n=9
        #self.fullText="DOCUMENTO NACIONAL DE IDENTIDAD NUMERO 12345678"
        for nameKey in nameKeys:
            if nameKey!="":
                reguexStr=r'(\d{%s})' % n
                pattern = r'%s%s%s' % (nameKey,midpattern, reguexStr)
                dnis=re.findall(pattern,self.fullText)
                if dnis!=[]:
                    for d in dnis:
                        dnislist.append(d)
        return dnislist
    def getdataword(self):
        dniKeys=self.configData.df["NRO DNI"].values.tolist()
        rucKeys=self.configData.df["RUC"].values.tolist()
        ceKeys=self.configData.df["CE"].values.tolist()
        partidaKeys=self.configData.df["PARTIDA"].values.tolist()
        nationalityKeys=self.configData.df["NACIONALIDAD"].values.tolist()
        sombrearKeys=self.configData.df["SOMBREAR"].values.tolist()
        correlativeKeys=self.configData.df["CORRELATIVOS"].values.tolist()
        self.getText()
        dnis=self.get_names("DNIS",dniKeys)
        dnis=list(dict.fromkeys(dnis))
        rucs=self.get_names("RUCS",rucKeys)
        rucs=list(dict.fromkeys(rucs))
        partidaE=self.get_names("PARTIDAS",partidaKeys)
        partidaE=list(dict.fromkeys(partidaE))
        carnets=self.get_names("CES",ceKeys)
        carnets=list(dict.fromkeys(carnets))
        #remove duplicates in list
        self.dataWord={
            "kardex":self.kardexName,
            "dnis":dnis,
            "rucs":rucs,
            "partidaE":partidaE,
            "carnets":carnets
        }
        return self.dataWord

def GettingkardexData():
    #get all .docx files in Kardexs folder
    pm=pathsManager()
    pathy=os.path.join(pm.currentFolderPath,"Kardexs")
    #print(pathy)
    files=[f for f in os.listdir(pathy) if f.endswith('.docx')]
    #print(files)
    listDataWord=[]
    for file in files:
        dWord=dataWord(file)
        #print(file)
        # llamar a mi funcion de busqueda de nombres naturales o juridicos
        listDataWord.append(dWord.getdataword())
        #print(dWord.getdataword())
    dataofWords={}

    for kardexData in listDataWord:
        #print(f"-------reading karex: {kardexData['kardex']}")
        dataofWords[kardexData['kardex']]={}
        dnilist=[]
        for dni in kardexData['dnis']:
            dnilist.append(dni)
        dataofWords[kardexData['kardex']]['DNI']=dnilist
        ruclist=[]
        for ruc in kardexData['rucs']:
            ruclist.append(ruc)
        dataofWords[kardexData['kardex']]['RUC']=ruclist
        partidList=[]
        for partid in kardexData['partidaE']:
            partidList.append(partid)
        dataofWords[kardexData['kardex']]['partidE']=partidList
    infoWords=os.path.join(pm.currentFolderPath,'dataofWords.json')
    with open(infoWords, 'w') as f:
        json.dump(dataofWords, f,indent=4)
    #print(listDataWord)
#sub_main()