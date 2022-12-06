from docx.enum.text import WD_COLOR_INDEX
from docx import Document
document = Document("dani.docx")
print(f"catindad de parrafos: {len(document.paragraphs)}")
searchedWord="Some"
for paragraph in document.paragraphs:
    if searchedWord in paragraph.text:
        #print(paragraph.text)
        #paragraph.font.highlight_color = WD_COLOR_INDEX.YELLOW
        for run in paragraph.runs:
            print(run.text)
            if searchedWord in run.text:
                x = run.text.split(searchedWord)
                print(f"cantidad de palabras en el run: {len(x)}")
                print(x)
                run.clear()

                for i in range(len(x)-1):
                    run.add_text(x[i])
                    paragraph.add_run(searchedWord).font.highlight_color = WD_COLOR_INDEX.YELLOW
                    #run.add_text(searchedWord)
                    #run.font.highlight_color = WD_COLOR_INDEX.YELLOW
document.save('dani2.docx')